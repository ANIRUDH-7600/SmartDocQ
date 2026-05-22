from __future__ import annotations

import io
import os
import sys
import types

import pytest


# Ensure backend/config.py can import in test environments.
os.environ.setdefault("SERVICE_TOKEN", "test")


# ---------------------------------------------------------------------------
# Stub external deps BEFORE importing indexer / retrieval_service
# ---------------------------------------------------------------------------


class _ImportCollectionStub:
    def upsert(self, *args, **kwargs):
        pass

    def get(self, *args, **kwargs):
        return {"ids": [], "documents": [], "metadatas": []}

    def delete(self, *args, **kwargs):
        pass

    def query(self, *args, **kwargs):
        return {"documents": [[]], "distances": [[]], "metadatas": [[]]}


fake_db_chroma = types.ModuleType("db.chroma")
fake_db_chroma.collection = _ImportCollectionStub()
sys.modules["db.chroma"] = fake_db_chroma


fake_embedding_service = types.ModuleType("services.embedding_service")
fake_embedding_service.generate_embeddings = lambda *_a, **_k: [0.0]
sys.modules["services.embedding_service"] = fake_embedding_service


from indexing import indexer  # noqa: E402
from services import retrieval_service  # noqa: E402

from utils.extraction import extract_text_from_docx_bytes  # noqa: E402
from utils.table_extraction import (  # noqa: E402
    extract_tables_from_csv_bytes,
    extract_tables_from_docx_bytes,
    extract_tables_from_xlsx_bytes,
    flatten_table_for_embedding,
    render_markdown_table,
)


# =========================
# Retrieval keyword extraction
# =========================


def _overlap(query: str, doc: str) -> int:
    return len(retrieval_service._keywords(query) & retrieval_service._keywords(doc))


def test_keywords_preserve_numeric_tokens():
    terms = retrieval_service._keywords("team 6 project title")
    assert "6" in terms
    assert _overlap("team 6 project title", "Team 6 final project submission") >= 3


def test_keywords_normalize_leading_zero_numbers():
    terms = retrieval_service._keywords("invoice 00045 approved")
    assert "45" in terms
    assert "00045" not in terms
    assert _overlap("invoice 00045", "Invoice 45 approved") >= 2


def test_keywords_do_not_normalize_alphanumeric_identifiers():
    terms = retrieval_service._keywords("47QZ9K2M7P CS001 A045X 00123AB")
    assert "47qz9k2m7p" in terms
    assert "cs001" in terms
    assert "a045x" in terms
    assert "00123ab" in terms


def test_keywords_support_team_number_overlap():
    assert _overlap("team 06 project", "Team 6 final project submission") >= 2


def test_keywords_support_marksheet_row_overlap():
    assert _overlap("marksheet row 003", "Row 3 marksheet for student") >= 3


def test_keywords_support_csv_xlsx_identifier_overlap():
    assert _overlap("Sheet1 row 06 total", "Sheet1 Row 6 Total = 95") >= 3


def test_keywords_filters_noise_and_stopwords():
    terms = retrieval_service._keywords("a, the; -- x 1 ?")
    assert terms == {"1"}


# =========================
# Table extraction
# =========================


def test_markdown_generation_basic():
    md = render_markdown_table(
        ["Name", "Marks", "Grade"],
        [["Rohit", "95", "A"], ["Anil", "82", "B"]],
    )
    assert "| Name" in md
    assert "| -----" in md
    assert "| Rohit" in md


def test_flattened_generation_basic():
    flat = flatten_table_for_embedding(
        sheet="Students",
        headers=["Name", "Marks", "Grade"],
        rows=[["Rohit", "95", "A"], ["Anil", "82", "B"]],
    )
    assert "Sheet: Students" in flat
    assert "Row:" in flat
    assert "Marks = 95" in flat


def test_csv_extraction_headers_and_rows():
    data = b"Name,Marks,Grade\nRohit,95,A\nAnil,82,B\n"
    tables = extract_tables_from_csv_bytes(data, sheet_name="Sheet1", source_key="doc1")
    assert len(tables) == 1

    t = tables[0]
    assert t["sheet"] == "Sheet1"
    assert t["headers"] == ["Name", "Marks", "Grade"]
    assert t["rows"][0] == ["Rohit", "95", "A"]
    assert "| Name" in t["markdown"]
    assert "Marks = 95" in t["flattened_text"]


def test_xlsx_extraction_multiple_sheets():
    try:
        from openpyxl import Workbook
    except Exception as e:  # pragma: no cover
        pytest.skip(f"openpyxl not available: {e}")

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Students"
    ws1.append(["Name", "Marks", "Grade"])
    ws1.append(["Rohit", 95, "A"])
    ws1.append(["Anil", 82, "B"])

    ws2 = wb.create_sheet("Totals")
    ws2.append(["Metric", "Value"])
    ws2.append(["Average", 88.5])

    bio = io.BytesIO()
    wb.save(bio)

    tables = extract_tables_from_xlsx_bytes(bio.getvalue(), source_key="doc2")
    sheets = {t["sheet"] for t in tables}
    assert "Students" in sheets
    assert "Totals" in sheets


def test_docx_table_extraction_order():
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()

    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Name"
    t1.rows[0].cells[1].text = "Marks"
    t1.rows[1].cells[0].text = "Rohit"
    t1.rows[1].cells[1].text = "95"

    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "Metric"
    t2.rows[0].cells[1].text = "Value"
    t2.rows[1].cells[0].text = "Total"
    t2.rows[1].cells[1].text = "177"

    bio = io.BytesIO()
    doc.save(bio)

    tables = extract_tables_from_docx_bytes(bio.getvalue(), source_key="doc3")
    assert len(tables) == 2

    assert tables[0]["headers"][:2] == ["Name", "Marks"]
    assert tables[1]["headers"][:2] == ["Metric", "Value"]


def test_docx_paragraph_extraction_normalizes_whitespace():
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()
    doc.add_paragraph("   Hello     world   ")
    doc.add_paragraph("   ")
    doc.add_paragraph("Line\t\twith\ninternal   whitespace")

    bio = io.BytesIO()
    doc.save(bio)

    text = extract_text_from_docx_bytes(bio.getvalue())
    assert "Hello world" in text
    assert "internal whitespace" in text
    # Ensure empty paragraphs do not introduce empty lines.
    assert "\n\n\n" not in text


def test_large_csv_extraction_completes_reasonably():
    import time

    rows = ["Name,Marks,Grade"]
    for i in range(10000):
        rows.append(f"Student{i},95,A")

    data = ("\n".join(rows)).encode("utf-8")

    start = time.perf_counter()

    tables = extract_tables_from_csv_bytes(
        data,
        sheet_name="Sheet1",
        source_key="perf_doc",
    )

    elapsed = time.perf_counter() - start

    assert len(tables) == 1
    assert len(tables[0]["rows"]) == 10000

    # Loose sanity threshold only (avoid brittle microbenchmarks).
    assert elapsed < 15


def test_corrupted_xlsx_returns_empty_tables():
    tables = extract_tables_from_xlsx_bytes(
        b"corrupted content",
        source_key="bad_doc",
    )

    assert tables == []


# =========================
# Hybrid indexing (text + tables)
# =========================


class FakeIndexerCollection:
    def __init__(self):
        self.store = {}

    def upsert(self, embeddings, documents, metadatas, ids):
        for emb, doc, meta, _id in zip(embeddings, documents, metadatas, ids):
            self.store[_id] = {"embedding": emb, "document": doc, "metadata": meta}

    def get(self, where=None, ids=None):
        if where and "doc_id" in where:
            doc_id = where["doc_id"]
            result_ids, result_docs, result_metas = [], [], []
            for _id, item in self.store.items():
                if item["metadata"].get("doc_id") == doc_id:
                    result_ids.append(_id)
                    result_docs.append(item["document"])
                    result_metas.append(item["metadata"])
            return {"ids": result_ids, "documents": result_docs, "metadatas": result_metas}
        return {"ids": list(self.store.keys()), "documents": [], "metadatas": []}

    def delete(self, ids=None, **_kwargs):
        for _id in (ids or []):
            self.store.pop(_id, None)


@pytest.fixture
def fake_indexer_collection(monkeypatch):
    coll = FakeIndexerCollection()
    monkeypatch.setattr(indexer, "collection", coll)
    return coll


@pytest.fixture(autouse=True)
def _no_node_push(monkeypatch):
    monkeypatch.setattr(indexer, "_push_chunks_to_node", lambda *_a, **_k: None)


@pytest.fixture(autouse=True)
def _mock_indexer_embeddings(monkeypatch):
    monkeypatch.setattr(indexer, "generate_embeddings", lambda _t: [0.0, 0.1, 0.2])


def test_csv_indexes_table_chunks_with_metadata(fake_indexer_collection):
    data = b"Name,Marks,Grade\nRohit,95,A\nAnil,82,B\n"

    ok, added = indexer.index_bytes(
        doc_id="doc_csv",
        filename="students.csv",
        mimetype="text/csv",
        data=data,
        file_hash="h",
    )

    assert ok is True
    assert added >= 1

    metas = [v["metadata"] for v in fake_indexer_collection.store.values()]
    table_metas = [m for m in metas if m.get("is_table")]
    assert table_metas, "expected at least one table chunk"

    m0 = table_metas[0]
    assert m0.get("table_id")
    assert m0.get("table_index") == 0
    assert m0.get("doc_id") == "doc_csv"
    assert m0.get("row_start") == 0
    assert m0.get("row_end") == 2

    # Flattened semantic storage is primary; markdown is preserved in metadata.
    any_table_id = next(
        _id for _id, v in fake_indexer_collection.store.items() if v["metadata"].get("is_table")
    )
    stored_doc = fake_indexer_collection.store[any_table_id]["document"]
    stored_meta = fake_indexer_collection.store[any_table_id]["metadata"]
    assert "Row:" in stored_doc
    assert "|" not in stored_doc
    assert "markdown" in stored_meta
    assert "|" in stored_meta.get("markdown", "")


def test_table_deduplication_skips_duplicate_docx_tables(fake_indexer_collection):
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()
    for _ in range(2):
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "Name"
        t.rows[0].cells[1].text = "Marks"
        t.rows[1].cells[0].text = "Rohit"
        t.rows[1].cells[1].text = "95"

    bio = io.BytesIO()
    doc.save(bio)

    ok, added = indexer.index_bytes(
        doc_id="doc_docx_dupe",
        filename="dupe_tables.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=bio.getvalue(),
        file_hash="h3",
    )

    assert ok is True
    # Two identical tables should dedup to one indexed table chunk.
    table_chunks = [v for v in fake_indexer_collection.store.values() if v["metadata"].get("is_table")]
    assert len(table_chunks) == 1


def test_chunk_id_uniqueness_uses_next_chunk_index(fake_indexer_collection, monkeypatch):
    """Regression test for chunk-id collisions when added != next_chunk_index.

    We simulate a divergence by monkeypatching _flush_batch to return 0 while
    still writing vectors to the collection.
    """

    # Force two text chunks
    monkeypatch.setattr(indexer, "split_sheet_sections", lambda _t: [(None, "body")])
    monkeypatch.setattr(indexer, "chunk_text", lambda _b: ["A" * 400, "B" * 400])

    # Fake a flush_batch that writes but reports 0 added
    real_upsert = fake_indexer_collection.upsert

    def fake_flush_batch(collection_ref, batch_embeddings, batch_documents, batch_metadatas, batch_ids):
        real_upsert(batch_embeddings, batch_documents, batch_metadatas, batch_ids)
        return 0

    monkeypatch.setattr(indexer, "_flush_batch", fake_flush_batch)

    # Provide one table via extractor
    monkeypatch.setattr(
        indexer,
        "extract_tables_for_file",
        lambda *_a, **_k: [
            {
                "table_id": "t1",
                "sheet": None,
                "headers": ["H1"],
                "rows": [["V1"]],
            }
        ],
    )

    ok, added = indexer.index_bytes(
        doc_id="doc_collision",
        filename="x.txt",
        mimetype="text/plain",
        data=b"body",
        file_hash="h4",
    )

    assert ok is True
    # We should have 3 unique ids: doc_collision_0, _1 (text) and _2 (table)
    assert set(fake_indexer_collection.store.keys()) == {"doc_collision_0", "doc_collision_1", "doc_collision_2"}


def test_deterministic_chunk_numbering_consumes_indices_on_embedding_fail(fake_indexer_collection, monkeypatch):
    # Force two text chunks in a deterministic order.
    monkeypatch.setattr(indexer, "split_sheet_sections", lambda _t: [(None, "body")])
    monkeypatch.setattr(indexer, "chunk_text", lambda _b: ["A" * 400, "B" * 400])

    calls = {"n": 0}

    def flaky_embeddings(_t):
        calls["n"] += 1
        # First chunk fails, second succeeds.
        return None if calls["n"] == 1 else [0.0, 0.1, 0.2]

    monkeypatch.setattr(indexer, "generate_embeddings", flaky_embeddings)

    ok, added = indexer.index_text("doc_det", "sample.txt", "ignored")
    assert ok is True
    assert added == 1

    # Chunk 0 was reserved but not stored; second chunk must be stored as chunk 1.
    assert set(fake_indexer_collection.store.keys()) == {"doc_det_1"}
    meta = fake_indexer_collection.store["doc_det_1"]["metadata"]
    assert meta.get("chunk") == 1


def test_markdown_truncation_in_metadata(fake_indexer_collection, monkeypatch):
    # Force markdown to exceed the metadata cap while keeping the table tiny.
    monkeypatch.setattr(
        indexer,
        "render_markdown_table",
        lambda *_a, **_k: ("| H |\n| --- |\n| V |\n" + ("x" * (indexer.MAX_MD_META_LEN + 200))),
    )
    monkeypatch.setattr(
        indexer,
        "flatten_table_for_embedding",
        lambda *_a, **_k: "Sheet: S1\n\nRow:\nH = V",
    )

    headers = ["H"]
    rows = [["V"]]
    tables = [{"table_id": "t_big", "sheet": "S1", "headers": headers, "rows": rows}]

    chunk_records = []
    added = indexer._index_tables(
        "doc_md",
        "big.xlsx",
        tables,
        start_chunk_index=0,
        chunk_records_out=chunk_records,
        file_hash="h_md",
    )

    assert added == 1
    stored = fake_indexer_collection.store["doc_md_0"]
    md_meta = stored["metadata"].get("markdown")
    assert isinstance(md_meta, str)
    assert len(md_meta) <= indexer.MAX_MD_META_LEN
    assert md_meta.endswith("...[truncated]")


def test_docx_table_only_is_indexed(fake_indexer_collection):
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        pytest.skip(f"python-docx not available: {e}")

    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"
    t.rows[0].cells[1].text = "Marks"
    t.rows[1].cells[0].text = "Rohit"
    t.rows[1].cells[1].text = "95"

    bio = io.BytesIO()
    doc.save(bio)

    ok, added = indexer.index_bytes(
        doc_id="doc_docx",
        filename="only_table.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=bio.getvalue(),
        file_hash="h2",
    )

    assert ok is True
    assert added >= 1

    metas = [v["metadata"] for v in fake_indexer_collection.store.values()]
    assert any(m.get("is_table") for m in metas)


# =========================
# Retrieval ranking
# =========================


class FakeRetrievalCollection:
    def __init__(self, docs, dists, metas):
        self._docs = docs
        self._dists = dists
        self._metas = metas

    def query(self, *args, **kwargs):
        return {
            "documents": [self._docs],
            "distances": [self._dists],
            "metadatas": [self._metas],
        }


@pytest.fixture(autouse=True)
def _mock_retrieval_embeddings(monkeypatch):
    monkeypatch.setattr(retrieval_service, "generate_embeddings", lambda _q: [0.0, 0.1, 0.2])


def test_table_question_boosts_table_chunks(monkeypatch):
    table_doc = "Sheet: Students\n\nRow:\nName = Rohit\nMarks = 95\nGrade = A"
    non_table_doc = ("This document explains grading policies and rules." * 2).strip()

    docs = [non_table_doc, table_doc]
    dists = [0.20, 0.25]  # non-table slightly closer
    metas = [
        {"is_table": False},
        {"is_table": True, "table_id": "t1", "table_index": 0},
    ]

    monkeypatch.setattr(
        retrieval_service,
        "collection",
        FakeRetrievalCollection(docs, dists, metas),
    )

    ctx, err = retrieval_service.retrieve_context(
        "What is the highest marks in the table?",
        "doc1",
    )

    assert err is None
    assert ctx is not None
    assert ctx.strip().startswith("Sheet: Students")


def test_generic_question_does_not_unfairly_boost_tables(monkeypatch):
    table_doc = "Sheet: Students\n\nRow:\nName = Rohit\nMarks = 95\nGrade = A"
    non_table_doc = ("This document explains grading policies and rules." * 2).strip()

    docs = [non_table_doc, table_doc]
    dists = [0.10, 0.20]  # non-table is closer, and query is generic.
    metas = [
        {"is_table": False},
        {"is_table": True, "table_id": "t1", "table_index": 0},
    ]

    monkeypatch.setattr(
        retrieval_service,
        "collection",
        FakeRetrievalCollection(docs, dists, metas),
    )

    ctx, err = retrieval_service.retrieve_context(
        "Explain the grading policies and rules",
        "doc1",
    )

    assert err is None
    assert ctx is not None
    # With no table intent, ranking should rely on similarity/overlap, not table boosts.
    assert ctx.strip().startswith("This document explains grading policies")
